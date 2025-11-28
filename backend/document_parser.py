"""
Document parser module for extracting text from PDF, DOCX, and TXT files.
"""
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
import charset_normalizer
from typing import Dict, List, Tuple
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse documents and extract text with metadata."""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.txt'}
    
    def parse(self, file_path: str) -> Dict:
        """
        Parse a document and extract text with metadata.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing:
                - text: Full extracted text
                - pages: List of page texts (for PDFs)
                - metadata: Document metadata
                - format: File format
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = path.suffix.lower()
        
        if suffix not in self.supported_formats:
            raise ValueError(f"Unsupported format: {suffix}")
        
        try:
            if suffix == '.pdf':
                return self._parse_pdf(file_path)
            elif suffix == '.docx':
                return self._parse_docx(file_path)
            elif suffix == '.txt':
                return self._parse_txt(file_path)
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {str(e)}")
            raise
    
    def _parse_pdf(self, file_path: str) -> Dict:
        """Parse PDF file using PyMuPDF."""
        doc = fitz.open(file_path)
        pages = []
        full_text = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            pages.append({
                'page_number': page_num + 1,
                'text': text
            })
            full_text.append(text)
        
        metadata = doc.metadata
        doc.close()
        
        return {
            'text': '\n\n'.join(full_text),
            'pages': pages,
            'metadata': metadata or {},
            'format': 'pdf',
            'total_pages': len(pages)
        }
    
    def _parse_docx(self, file_path: str) -> Dict:
        """Parse DOCX file using python-docx."""
        doc = Document(file_path)
        paragraphs = []
        full_text = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
                full_text.append(text)
        
        # Extract tables
        tables_text = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(' | '.join(row_data))
            tables_text.append('\n'.join(table_data))
        
        if tables_text:
            full_text.extend(tables_text)
        
        metadata = {
            'paragraphs': len(paragraphs),
            'tables': len(doc.tables)
        }
        
        return {
            'text': '\n\n'.join(full_text),
            'pages': [{'page_number': 1, 'text': '\n\n'.join(full_text)}],
            'metadata': metadata,
            'format': 'docx',
            'total_pages': 1  # DOCX doesn't have page concept
        }
    
    def _parse_txt(self, file_path: str) -> Dict:
        """Parse TXT file with encoding detection."""
        # Read raw bytes
        with open(file_path, 'rb') as f:
            raw_data = f.read()
        
        # Detect encoding
        detection = charset_normalizer.from_bytes(raw_data).best()
        
        if detection is None:
            # Fallback to UTF-8
            text = raw_data.decode('utf-8', errors='ignore')
            encoding = 'utf-8'
        else:
            text = str(detection)
            encoding = detection.encoding
        
        return {
            'text': text,
            'pages': [{'page_number': 1, 'text': text}],
            'metadata': {'encoding': encoding},
            'format': 'txt',
            'total_pages': 1
        }
